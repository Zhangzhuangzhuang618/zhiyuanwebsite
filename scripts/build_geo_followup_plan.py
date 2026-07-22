from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


OUTPUT = Path(__file__).resolve().parents[1] / "志远搬家GEO优化方案（后续推进版）.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GOLD = "7A5A00"
RED = "9B1C1C"
FONT = "Arial Unicode MS"


def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_run(p, text, **kwargs):
    run = p.add_run(text)
    set_font(run, **kwargs)
    return run


def add_body(doc, text, bold_prefix=None, color=INK, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, size=11, color=color, bold=True)
        add_run(p, text[len(bold_prefix):], size=11, color=color)
    else:
        add_run(p, text, size=11, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    add_run(p, text, size=10.5, color=INK)
    return p


def add_heading(doc, level, text):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    add_run(p, text, size={1: 16, 2: 13, 3: 12}[level], color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    return p


def add_callout(doc, label, text, tone="blue"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE if tone == "blue" else LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, label + "  ", size=10.5, color=DARK_BLUE if tone == "blue" else GOLD, bold=True)
    add_run(p, text, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, header, size=10, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            p = cells[i].paragraphs[0]
            if i in (0, 1) and len(text) < 13:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for level, size, color, before, after in ((1, 16, BLUE, 16, 8), (2, 13, BLUE, 12, 6), (3, 12, DARK_BLUE, 8, 4)):
        style = styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    setup_styles(doc)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header_p, "志远搬家｜GEO 后续推进", size=9, color=MUTED)
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer_p, "内部执行文件｜v2.0｜2026年7月", size=9, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after = Pt(12)
    add_run(p, "志远搬家 GEO 优化方案", size=25, color=INK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    add_run(p, "后续推进版｜仅保留尚未完成的工作", size=14, color=DARK_BLUE)
    add_callout(doc, "适用范围", "面向中国大陆搜索与生成式搜索场景，重点覆盖广州及周边真实服务区域。本文不重复已完成的站内基础技术改造。")
    add_body(doc, "目标网站：www.zhiyuanbj.cn", bold_prefix="目标网站：", after=3)
    add_body(doc, "版本：v2.0｜更新日期：2026年7月15日", bold_prefix="版本：", after=18)
    add_body(doc, "执行目标：在不虚构资质、案例、评价、价格或服务网点的前提下，补足本地实体可信度、区域内容、第三方佐证、收录验证与持续运营能力。", bold_prefix="执行目标：", after=6)
    add_body(doc, "评分口径：当前约为 66/100 的源码 GEO 完成度；本方案的目标是将可验证完成度提升至约 80/100。该分数不等同于百度排名、AI 引用次数或获客结果。", bold_prefix="评分口径：", after=18)

    add_heading(doc, 1, "一、后续工作边界与原则")
    add_bullet(doc, "只建设真实可核验的本地信息。公司名称、地址、电话、营业时间、服务区域、资质、案例、价格和评价必须有内部凭证或公开依据。")
    add_bullet(doc, "区域页必须对应实际服务能力与独立素材。不能仅替换“广州/天河/佛山”等地名批量生成近似页面。")
    add_bullet(doc, "第三方平台仅做正式认领、真实资料完善和自然内容发布；不购买虚假评价、不伪造媒体报道、不制造百科词条。")
    add_bullet(doc, "所有优化以收录、抓取、内容质量和有效咨询数据复盘；不承诺特定关键词排名或 AI 回答展示。")

    add_heading(doc, 1, "二、分阶段推进路线")
    add_table(doc,
        ["阶段", "建议周期", "核心交付", "验收标准"],
        [
            ("0", "第 1 周", "线上部署与收录基线", "公开可访问；完成平台验证与数据建档"),
            ("1", "第 1—2 周", "本地实体资料与信任素材", "资料可核验、官网与平台口径一致"),
            ("2", "第 2—6 周", "广州及周边差异化内容", "每页有独立事实、场景与服务边界"),
            ("3", "持续推进", "第三方实体与内容佐证", "完成认领/发布，保留链接与后台截图"),
            ("4", "每月/每季", "性能、抓取、内容复盘", "形成问题清单、修复记录与更新日志"),
        ],
        [720, 1080, 3960, 3600],
    )

    add_heading(doc, 1, "三、阶段 0：上线、抓取与收录验证")
    add_body(doc, "目的：确认已经完成的技术基础真正部署在生产站，而不是仅存在于本地源码。")
    add_heading(doc, 2, "3.1 生产环境核验")
    for text in [
        "逐页检查首页、服务页、案例页、新闻详情、FAQ、联系页的 HTTP 状态码、canonical、标题、描述、移动端渲染与页面主内容是否正常。",
        "确认 robots.txt、llms.txt 与 sitemap.xml 能通过正式域名访问；检查 sitemap 中 URL 的域名、更新时间和返回状态，剔除跳转、重复或失效页面。",
        "确认 HTTPS 证书、www/非 www 跳转规则、404 页面和旧 URL 跳转策略一致，避免同一内容被多个地址重复收录。",
        "从服务器日志或 CDN 日志识别百度、搜狗、360及允许的检索爬虫访问情况；记录抓取频率、404、5xx 与被拒绝 URL。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, 2, "3.2 中国大陆搜索平台接入")
    for text in [
        "在百度搜索资源平台验证正式域名，提交 sitemap.xml，并持续查看抓取异常、索引覆盖、死链和移动适配报告。",
        "视业务流量情况接入 360、搜狗等搜索站长平台；Google Search Console 作为补充，不应替代百度侧验证。",
        "建立首份基线表：已提交 URL 数、已收录 URL 数、品牌词与核心服务词展示/点击、主要落地页、有效咨询。后续按月对比趋势。",
    ]:
        add_bullet(doc, text)
    add_callout(doc, "验收材料", "平台验证截图、sitemap 提交记录、首轮抓取/索引报表、核心页面检查清单。")

    add_heading(doc, 1, "四、阶段 1：本地实体与可信信息补齐")
    add_body(doc, "目的：让官网、地图与第三方页面对“这是谁、在哪里、提供什么服务”给出一致、可验证的答案。")
    add_heading(doc, 2, "4.1 资料台账（先收集，后发布）")
    add_table(doc,
        ["资料类别", "需要确认的事实", "发布位置"],
        [
            ("主体信息", "营业执照主体名称、统一社会信用代码、品牌名称", "官网关于页、地图/商户认领资料"),
            ("联系与地址", "真实服务地址、电话、营业时间、客服职责", "联系页、页脚、地图平台"),
            ("服务范围", "实际覆盖的广州区县及周边城市、可提供的服务类型", "服务页、区域页、FAQ"),
            ("服务依据", "合同、报价规则、赔付/保险条件、设备与人员能力", "服务说明、常见问题、合同流程"),
            ("经验素材", "获授权的案例、现场照片、流程记录、客户评价", "案例页、图文/视频内容"),
        ],
        [1560, 4080, 3720],
    )
    add_heading(doc, 2, "4.2 官网信任内容")
    for text in [
        "完善“关于志远/联系志远”页面：明确主体、服务定位、真实服务区域、电话、营业时间及投诉/售后联系路径。",
        "为可公开的服务承诺增加条件与边界，例如报价确认方式、临时增项处理、物品损坏申报流程。没有书面规则的内容不发布为承诺。",
        "建立案例模板：服务日期（可只到月份）、起止区域、搬迁类型、物品规模、使用车型/人员、难点与处理方式、客户授权状态。避免展示客户隐私。",
        "为文章和指南增加作者/审核人、发布日期、最近更新日期及资料来源；只有具备真实责任主体时才展示。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, 1, "五、阶段 2：广州及周边差异化内容建设")
    add_body(doc, "目的：从泛化搬家介绍，升级为能回答本地用户具体问题的内容体系。优先做高意图服务页和真实覆盖区域，数量服从质量。")
    add_heading(doc, 2, "5.1 区域页建设规则")
    for text in [
        "先由业务确认实际常态化服务的区域，再决定是否建设页面。广州各区、佛山、东莞、中山等仅在确有服务能力、接单规则和可用素材时纳入。",
        "每个区域页至少包含：该区域服务边界、常见住宅/写字楼搬运条件、停车/电梯/进场限制提示、适用服务、真实案例或现场素材、区域常见问题与联系方式。",
        "若区域暂时没有独立事实和案例，保留在“服务范围”汇总页，不建立独立落地页，防止重复内容和站群风险。",
        "区域页标题、描述、正文、图片、FAQ 与案例必须独立撰写；不要只替换地名或复制主站模板。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, 2, "5.2 服务与指南内容优先级")
    add_table(doc,
        ["优先内容", "应回答的问题", "所需事实材料"],
        [
            ("同城家庭搬家", "车型/人员如何匹配？楼层、电梯、停车如何影响安排？", "服务流程、实际限制、车型/人员配置规则"),
            ("跨市搬迁", "交接、装卸、运输与到达如何安排？", "真实覆盖线路、时效边界、交接流程"),
            ("企业/办公室搬迁", "IT、档案、工位、停业窗口如何组织？", "项目流程、负责人、授权案例"),
            ("特殊物品与高空吊装", "哪些物品/楼宇条件适用？风险如何评估？", "设备能力、安全条件、禁止情形"),
            ("搬家指南", "用户要如何打包、断电、交接与验收？", "可复核的流程、清单、编辑/审核人"),
        ],
        [1920, 3600, 3840],
    )
    add_heading(doc, 2, "5.3 FAQ 与内容更新")
    for text in [
        "把现有 FAQ 按搜索意图补充为本地化问题：区域停车、电梯、城中村/小区进场规则、跨市交接、办公搬迁、特殊物品等。每个答案先给结论，再说明条件。",
        "涉及价格、时效、服务覆盖、赔付或资质时，只写已确认且有日期的事实；价格应说明影响因素和询价方式，不发布无法长期保证的区间。",
        "每季度复核 FAQ、服务页与指南中的价格、区域、联系方式、流程和案例；在页面展示最近更新日期，并留存内部修订记录。",
        "对仍为空的内容图片进行全量 alt 审计：仅为实际表达内容的图片补充准确描述；装饰图标保持空 alt，避免堆砌关键词。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, 1, "六、阶段 3：第三方实体与权威佐证")
    add_body(doc, "目的：建立独立来源对企业实体和服务能力的交叉印证。此阶段需要运营人员或企业负责人提供资料并完成账号认领。")
    add_table(doc,
        ["渠道", "应做事项", "边界与验收"],
        [
            ("百度地图 / 高德 / 腾讯地图", "认领或创建真实门店/服务点，统一主体、地址、电话、营业时间与服务类目。", "仅认领实际存在地点；保存审核结果与页面链接。"),
            ("大众点评及本地生活平台", "完善商户资料、服务说明、真实图片与售后联系。", "不购买或诱导虚假评价；评价必须来自真实服务。"),
            ("58同城等分类信息平台", "按平台规则完成企业认证和服务发布。", "信息与官网一致；不批量复制内容或发布虚假网点。"),
            ("媒体 / 行业组织 / 社区", "发布有事实依据的搬家避坑、流程、安全或行业观察内容。", "保留原始链接与发布依据；不伪造报道或付费软文事实。"),
            ("百科与问答平台", "仅在满足平台收录及可引用来源要求时申请词条或参与问答。", "不能将百科创建当作必达任务；回答需解决问题而非硬广。"),
            ("短视频与图文内容", "将真实案例、打包技巧、流程演示拆分为可检索主题。", "取得客户和员工出镜授权；每条内容可回链至对应官网页。"),
        ],
        [1800, 4020, 3540],
    )

    add_heading(doc, 1, "七、阶段 4：性能、结构化扩展与持续复盘")
    add_heading(doc, 2, "7.1 性能与可访问性")
    for text in [
        "使用真实中国大陆网络环境测试首页、服务页和报价页的移动端首屏速度、图片体积、缓存、渲染阻塞资源和验证码可用性；按测试结果逐项优化，不以“合并全部 CSS/JS”为前提。",
        "定期爬取正式站，修复 404、重定向链、重复标题/描述、孤立页、失效图片和 sitemap 与实际页面不一致的问题。",
        "维护报价表单、电话按钮和在线咨询的可用性；这些是本地服务站最直接的转化路径，应纳入每月检查。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, 2, "7.2 结构化数据的增量工作")
    for text in [
        "针对已发布且步骤明确的原创搬家指南，补充 HowTo 结构化数据；步骤、工具、时长等字段必须与页面正文一致。",
        "区域页仅在内容真实独立时增加对应的 service area 信息；不要虚构“分部”、地址或营业时间。",
        "新增案例和文章时持续校验 JSON-LD 的标题、日期、图片、作者与正文一致。结构化数据用于机器理解，不应替代真实内容。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, 2, "7.3 月度与季度复盘指标")
    add_table(doc,
        ["周期", "检查指标", "输出物"],
        [
            ("每月", "抓取/索引异常、有效收录页、品牌词与核心服务词展现、页面 404、表单/电话可用性", "月度监测表与修复清单"),
            ("每季", "FAQ/服务范围/联系方式/价格说明/案例时效、区域页重复度、第三方资料一致性", "内容更新记录与下季选题"),
            ("每半年", "真实客户咨询来源、重点区域需求、外部引用质量、站点性能", "优先级调整与资源投入建议"),
        ],
        [1320, 5040, 3000],
    )

    add_heading(doc, 1, "八、执行前需由企业提供的材料")
    add_callout(doc, "阻塞条件", "以下资料缺失时，不应通过猜测或营销措辞代替；应先补充凭证，再发布对应页面或第三方资料。", tone="gray")
    for text in [
        "营业执照主体信息、品牌使用授权（如主体与品牌不完全一致）、可公开的联系方式和服务地址。",
        "真实服务区域、可提供的服务类型、车型/人员/设备能力及高空吊装等特殊服务的适用条件。",
        "报价、合同、增项、赔付、保险和售后流程的现行书面规则。",
        "经客户授权的案例素材：照片/视频、时间、区域、服务类型、可公开的处理过程与评价。",
        "各平台既有账号、地图门店认领权限、百度站长平台/域名验证所需权限。",
        "内容审核责任人：业务审核、法务/负责人审核、发布人及季度复核责任人。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, 1, "九、优先级清单")
    add_table(doc,
        ["优先级", "下一步", "负责人/前置条件", "完成判定"],
        [
            ("P0", "生产环境与百度资源平台验证", "运维 + 域名/平台权限", "可访问、已提交、基线已记录"),
            ("P0", "整理真实实体与服务资料台账", "企业负责人 + 业务", "每项有来源、可公开范围明确"),
            ("P1", "广州真实服务区域的差异化内容", "内容 + 业务审核 + 案例素材", "非模板化、含独立事实与 FAQ"),
            ("P1", "地图与本地生活平台正式认领", "账号权限 + 实体证明", "资料一致、审核通过或留存原因"),
            ("P1", "服务页/案例/指南的可信内容", "业务资料 + 客户授权", "作者/审核/日期/证据可追溯"),
            ("P2", "性能与结构化数据增量优化", "线上测试数据", "按问题清单修复并复测"),
        ],
        [720, 3000, 3120, 2520],
    )
    add_body(doc, "结论：接下来的瓶颈不在继续堆叠标签，而在真实本地资料、区域差异化内容、第三方实体佐证和生产环境数据闭环。", bold_prefix="结论：", after=0)

    doc.core_properties.title = "志远搬家 GEO 优化方案（后续推进版）"
    doc.core_properties.subject = "中国大陆及广州周边本地服务 GEO 后续执行计划"
    doc.core_properties.author = "志远搬家"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
